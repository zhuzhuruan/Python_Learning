# -*- coding: utf-8 -*-
# @Time : 2023/8/112:20
# @Author : cyy
# @File : white_sample_deal.py
# @desc : 提取文档中的表格并对权限列进行清洗

import os
import docx
import pandas as pd
import re


def get_path_list(file):
    """
    获取文件列表
    :return:
    """
    file_list_1 = []
    file_list_2 = []
    for file_path, dir_name, file_names in os.walk(file):
        for file in file_names:
            file_name = os.path.join(file_path, file)
            if file_name.endswith('docx'):
                file_list_1.append(file_name)
            else:
                file_list_2.append(file_name)

    return file_list_1, file_list_2


def check_wait_deal_index(tables):
    """
    识别需要处理的table的序号
    :param tables:
    :return:
    """
    table_index = {}
    for i in range(len(tables)):
        row = tables[i].rows[0]
        column_list = [cell.text for cell in row.cells]
        if column_list == ['字段', '内容']:
            table_index['字段'] = i
        elif column_list == ['权限', '权限说明', '权限分类']:
            table_index['权限'] = i
    return table_index


def get_permission_table(path):
    """
    读取word文档中指定位置的表格并转换成dataframe
    :param path: .docx
    :param args: 3,5
    :return: dataframe
    """
    # 读取表格
    document = docx.Document(path)
    tables = document.tables
    # 获取需要读取的表格序号
    table_index = check_wait_deal_index(tables)
    name_info_table = tables[table_index['字段']]
    permission_table = tables[table_index['权限']]

    # 基本信息
    row_dic = {}
    for row in name_info_table.rows[1:]:
        row_dic[list(row.cells)[0].text] = list(row.cells)[1].text

    # 权限信息
    permission_list = []
    permission = []
    for row in permission_table.rows:
        cell_list = [cell.text for cell in row.cells]
        if cell_list[0] != '权限':
            permission_dic = {'permission': cell_list[0], 'description': cell_list[1], 'classification': cell_list[2]}
            permission_list.append(permission_dic)
            permission.append(cell_list[0])
    row_dic['permission'] = permission
    row_dic['permission_list'] = permission_list

    # 文件名：用来修正app的名称
    row_dic['file_name'] = re.search(r'.*?(?=\.)', os.path.basename(path)).group()
    return row_dic


def change_permission_output(df):
    """
    把权限及其说明单独输出
    :return:
    """
    df['permission_list'] = df['permission_list'].apply(lambda x: eval(x))
    x = sum(df['permission_list'].tolist(), [])
    permission_desc = pd.DataFrame(x)
    return permission_desc


def main_process(in_path):
    """
    读取多个文件处理后拼接
    :return:
    """
    file_list1, file_list2 = get_path_list(in_path)
    permission_list = []
    for file in file_list1:
        # print(file)
        row_dic = get_permission_table(file)
        permission_list.append(row_dic)
    permission_table = pd.DataFrame(permission_list)
    return permission_table


if __name__ == '__main__':
    base_path = r'C:\Users\caoyuanyuan\Desktop\项目文档\GA项目\3-模型\白样本'
    in_file = '解压'
    out_file = 'permission_table.xlsx'

    in_path = os.path.join(base_path, in_file)
    out_path = os.path.join(base_path, out_file)
    # 清洗权限
    # permission_table = main_process(in_path)
    # permission_table.to_excel(out_path)
    # # 获得权限说明
    # permission_desc = change_permission_output(permission_table)
    # permission_desc.to_excel(os.path.join(base_path, 'permission_desc.xlsx'))

    # path = r'C:\Users\caoyuanyuan\Desktop\项目文档\GA项目\3-模型\白样本\解压\交通旅游45-word\交通旅游45-word\12306智行火车票.apk.docx'
    # document = docx.Document(path)
    # tables = document.tables

    df = pd.read_excel(out_path)
    permission_desc = change_permission_output(df)
    permission_desc.to_excel(os.path.join(base_path, 'permission_desc.xlsx'))
